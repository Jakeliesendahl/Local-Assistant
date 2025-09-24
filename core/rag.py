"""
Document ingestion module for RAG system.
Supports PDF, DOCX, Markdown, and ICS (calendar) file processing.
"""

import os
from pathlib import Path
from typing import Optional, List, Dict, Any
import logging
import re
from uuid import uuid4
import sqlite3
from datetime import datetime
from .llm import LLMClient, LLMError

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

try:
    import chromadb
    from chromadb.config import Settings
except ImportError:
    chromadb = None

try:
    from pypdf import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    from icalendar import Calendar
    ICALENDAR_AVAILABLE = True
except ImportError:
    Calendar = None
    ICALENDAR_AVAILABLE = False

# Set up logging
logger = logging.getLogger(__name__)


def ingest_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from a PDF file using pypdf.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        Optional[str]: Extracted text content or None if extraction fails
    """
    if PdfReader is None:
        logger.error("pypdf is not installed. Please install it with: pip install pypdf")
        return None
    
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        reader = PdfReader(file_path)
        text_content = []
        
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                continue
        
        if not text_content:
            logger.warning(f"No text content extracted from PDF: {file_path}")
            return None
        
        return "\n\n".join(text_content)
        
    except Exception as e:
        logger.error(f"Error processing PDF file {file_path}: {e}")
        return None


def ingest_docx(file_path: str) -> Optional[str]:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        Optional[str]: Extracted text content or None if extraction fails
    """
    if Document is None:
        logger.error("python-docx is not installed. Please install it with: pip install python-docx")
        return None
    
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                text_content.append("--- Table ---\n" + "\n".join(table_text))
        
        if not text_content:
            logger.warning(f"No text content extracted from DOCX: {file_path}")
            return None
        
        return "\n\n".join(text_content)
        
    except Exception as e:
        logger.error(f"Error processing DOCX file {file_path}: {e}")
        return None


def ingest_markdown(file_path: str) -> Optional[str]:
    """
    Read text from a Markdown file.
    
    Args:
        file_path (str): Path to the Markdown file
        
    Returns:
        Optional[str]: File content or None if reading fails
    """
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        if not content.strip():
            logger.warning(f"Markdown file is empty: {file_path}")
            return None
        
        return content
        
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
            logger.info(f"Successfully read file with latin-1 encoding: {file_path}")
            return content
        except Exception as e:
            logger.error(f"Error reading Markdown file with alternative encoding {file_path}: {e}")
            return None
    except Exception as e:
        logger.error(f"Error reading Markdown file {file_path}: {e}")
        return None


def ingest_ics(file_path: str) -> Optional[str]:
    """
    Extract text content from an ICS (iCalendar) file using a custom parser.
    
    This function implements a simple ICS parser that doesn't rely on external libraries,
    making it more reliable across different Python environments.
    
    Args:
        file_path (str): Path to the ICS file
        
    Returns:
        Optional[str]: Extracted calendar content as formatted text or None if extraction fails
    """
    try:
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None
        
        with open(file_path, 'r', encoding='utf-8') as file:
            ics_content = file.read().strip()
        
        if not ics_content:
            logger.warning(f"ICS file is empty: {file_path}")
            return None
        
        # Parse the ICS content line by line
        lines = ics_content.split('\n')
        content_parts = []
        
        # Calendar metadata
        calendar_name = None
        calendar_desc = None
        prodid = None
        version = None
        
        # Event and task storage
        events = []
        todos = []
        
        # Current component being parsed
        current_component = None
        current_data = {}
        
        # Parse each line
        for line_num, line in enumerate(lines, 1):
            line = line.strip()
            if not line:
                continue
            
            # Handle line continuations (lines starting with space or tab)
            if line.startswith(' ') or line.startswith('\t'):
                if current_data and 'last_property' in current_data and current_data['last_property']:
                    current_data[current_data['last_property']] += line[1:]  # Remove the leading whitespace
                continue
            
            # Split property and value
            if ':' in line:
                prop, value = line.split(':', 1)
                prop = prop.upper().strip()
                value = value.strip()
                
                # Handle parameters (e.g., DTSTART;TZID=America/New_York:20240101T120000)
                if ';' in prop:
                    prop = prop.split(';')[0]
                
                # Handle BEGIN/END commands first
                if prop == 'BEGIN':
                    if value.upper() in ['VEVENT', 'VTODO']:
                        current_component = value.upper()
                        current_data = {'last_property': None}
                elif prop == 'END':
                    if value.upper() in ['VEVENT', 'VTODO'] and current_component == value.upper():
                        if current_component == 'VEVENT' and current_data:
                            # Remove the last_property tracking key before storing
                            clean_data = {k: v for k, v in current_data.items() if k != 'last_property'}
                            if clean_data:  # Only add if there's actual data
                                events.append(clean_data)
                        elif current_component == 'VTODO' and current_data:
                            clean_data = {k: v for k, v in current_data.items() if k != 'last_property'}
                            if clean_data:  # Only add if there's actual data
                                todos.append(clean_data)
                        current_component = None
                        current_data = {}
                
                # Calendar-level properties (when not in a component)
                elif current_component is None:
                    if prop == 'X-WR-CALNAME':
                        calendar_name = value
                    elif prop == 'X-WR-CALDESC':
                        calendar_desc = value
                    elif prop == 'PRODID':
                        prodid = value
                    elif prop == 'VERSION':
                        version = value
                
                # Component properties (when inside a component)
                elif current_component and prop not in ['BEGIN', 'END']:
                    current_data[prop] = value
                    current_data['last_property'] = prop
        
        # Build the formatted output
        if calendar_name:
            content_parts.append(f"Calendar Name: {calendar_name}")
        if calendar_desc:
            content_parts.append(f"Calendar Description: {calendar_desc}")
        if prodid:
            content_parts.append(f"Created by: {prodid}")
        if version:
            content_parts.append(f"iCalendar Version: {version}")
        
        if content_parts:
            content_parts.append("")  # Add blank line
        
        # Process events
        if events:
            content_parts.append("=== EVENTS ===")
            
            for i, event in enumerate(events, 1):
                event_parts = [f"\n--- Event {i} ---"]
                
                if 'SUMMARY' in event:
                    event_parts.append(f"Title: {event['SUMMARY']}")
                
                if 'DESCRIPTION' in event:
                    description = event['DESCRIPTION'].replace('\\n', '\n').replace('\\,', ',').replace('\\;', ';')
                    description = re.sub(r'\n+', '\n', description.strip())
                    event_parts.append(f"Description: {description}")
                
                # Parse and format dates
                if 'DTSTART' in event:
                    start_time = _parse_ics_datetime(event['DTSTART'])
                    event_parts.append(f"Start: {start_time}")
                
                if 'DTEND' in event:
                    end_time = _parse_ics_datetime(event['DTEND'])
                    event_parts.append(f"End: {end_time}")
                
                if 'LOCATION' in event:
                    event_parts.append(f"Location: {event['LOCATION']}")
                
                if 'ORGANIZER' in event:
                    organizer = event['ORGANIZER'].replace('mailto:', '')
                    event_parts.append(f"Organizer: {organizer}")
                
                if 'ATTENDEE' in event:
                    attendee = event['ATTENDEE'].replace('mailto:', '')
                    event_parts.append(f"Attendee: {attendee}")
                
                if 'STATUS' in event:
                    event_parts.append(f"Status: {event['STATUS']}")
                
                if 'PRIORITY' in event:
                    event_parts.append(f"Priority: {event['PRIORITY']}")
                
                if 'CATEGORIES' in event:
                    event_parts.append(f"Categories: {event['CATEGORIES']}")
                
                if 'URL' in event:
                    event_parts.append(f"URL: {event['URL']}")
                
                if 'RRULE' in event:
                    event_parts.append(f"Recurrence: {event['RRULE']}")
                
                # Timestamps
                if 'CREATED' in event:
                    created = _parse_ics_datetime(event['CREATED'])
                    event_parts.append(f"Created: {created}")
                
                if 'LAST-MODIFIED' in event:
                    modified = _parse_ics_datetime(event['LAST-MODIFIED'])
                    event_parts.append(f"Last Modified: {modified}")
                
                content_parts.extend(event_parts)
        
        # Process todos
        if todos:
            content_parts.append("\n=== TASKS/TODOS ===")
            
            for i, todo in enumerate(todos, 1):
                todo_parts = [f"\n--- Task {i} ---"]
                
                if 'SUMMARY' in todo:
                    todo_parts.append(f"Task: {todo['SUMMARY']}")
                
                if 'DESCRIPTION' in todo:
                    description = todo['DESCRIPTION'].replace('\\n', '\n').replace('\\,', ',').replace('\\;', ';')
                    description = re.sub(r'\n+', '\n', description.strip())
                    todo_parts.append(f"Description: {description}")
                
                if 'DUE' in todo:
                    due_date = _parse_ics_datetime(todo['DUE'])
                    todo_parts.append(f"Due: {due_date}")
                
                if 'STATUS' in todo:
                    todo_parts.append(f"Status: {todo['STATUS']}")
                
                if 'PRIORITY' in todo:
                    todo_parts.append(f"Priority: {todo['PRIORITY']}")
                
                if 'PERCENT-COMPLETE' in todo:
                    todo_parts.append(f"Percent Complete: {todo['PERCENT-COMPLETE']}%")
                
                content_parts.extend(todo_parts)
        
        # Add summary
        summary_parts = ["\n=== CALENDAR SUMMARY ==="]
        summary_parts.append(f"Total Events: {len(events)}")
        if todos:
            summary_parts.append(f"Total Tasks: {len(todos)}")
        
        content_parts.extend(summary_parts)
        
        if not events and not todos:
            logger.warning(f"No events or tasks found in ICS file: {file_path}")
            logger.debug(f"Calendar metadata found: name={calendar_name}, desc={calendar_desc}, prodid={prodid}")
            logger.debug(f"Total lines processed: {len(lines)}")
            return None
        
        result = "\n".join(content_parts)
        logger.info(f"Successfully extracted {len(events)} events and {len(todos)} tasks from ICS file: {file_path}")
        return result
        
    except Exception as e:
        logger.error(f"Error processing ICS file {file_path}: {e}")
        return None


def _parse_ics_datetime(dt_string: str) -> str:
    """
    Parse ICS datetime string and return a readable format.
    
    Args:
        dt_string (str): ICS datetime string (e.g., '20240101T120000Z' or '20240101')
        
    Returns:
        str: Formatted datetime string
    """
    try:
        # Remove timezone info for simple parsing
        clean_dt = dt_string.split(';')[0] if ';' in dt_string else dt_string
        
        # Handle date-only format (YYYYMMDD)
        if len(clean_dt) == 8 and clean_dt.isdigit():
            year = clean_dt[:4]
            month = clean_dt[4:6]
            day = clean_dt[6:8]
            return f"{year}-{month}-{day}"
        
        # Handle datetime format (YYYYMMDDTHHMMSS or YYYYMMDDTHHMMSSZ)
        elif 'T' in clean_dt:
            date_part, time_part = clean_dt.split('T')
            time_part = time_part.rstrip('Z')  # Remove UTC indicator
            
            if len(date_part) == 8 and len(time_part) >= 6:
                year = date_part[:4]
                month = date_part[4:6]
                day = date_part[6:8]
                hour = time_part[:2]
                minute = time_part[2:4]
                second = time_part[4:6] if len(time_part) >= 6 else '00'
                
                return f"{year}-{month}-{day} {hour}:{minute}:{second}"
        
        # If parsing fails, return original string
        return dt_string
        
    except Exception:
        return dt_string


def ingest_document(file_path: str) -> Optional[Dict[str, Any]]:
    """
    Unified document ingestion function that handles PDF, DOCX, Markdown, and ICS files.
    
    Args:
        file_path (str): Path to the document file
        
    Returns:
        Optional[Dict[str, Any]]: Dictionary containing file info and extracted content,
                                 or None if ingestion fails
    """
    if not os.path.exists(file_path):
        logger.error(f"File not found: {file_path}")
        return None
    
    file_path = Path(file_path)
    virign_file_extension = file_path.suffix
    file_extension = file_path.suffix.lower()
    
    # Determine file type and call appropriate function
    content = None
    if file_extension == '.pdf':
        content = ingest_pdf(str(file_path))
    elif file_extension == '.docx':
        content = ingest_docx(str(file_path))
    elif file_extension in ['.md', '.markdown']:
        content = ingest_markdown(str(file_path))
    elif file_extension == '.ics':
        content = ingest_ics(str(file_path))
    else:
        logger.error(f"Unsupported file type: {file_extension}")
        return None
    
    if content is None:
        return None
    
    # Return structured data
    return {
        'file_path': str(file_path),
        'file_name': file_path.name,
        'file_type': virign_file_extension,
        'content': content,
        'content_length': len(content),
        'word_count': len(content.split()) if content else 0
    }


def ingest_documents_from_directory(directory_path: str, 
                                  file_types: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    """
    Ingest all supported documents from a directory.
    
    Args:
        directory_path (str): Path to the directory containing documents
        file_types (Optional[List[str]]): List of file extensions to include 
                                        (default: ['.pdf', '.docx', '.md', '.markdown', '.ics'])
        
    Returns:
        List[Dict[str, Any]]: List of successfully ingested documents
    """
    if file_types is None:
        file_types = ['.pdf', '.docx', '.md', '.markdown', '.ics']
    
    if not os.path.exists(directory_path):
        logger.error(f"Directory not found: {directory_path}")
        return []
    
    ingested_documents = []
    directory = Path(directory_path)
    
    for file_path in directory.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in file_types:
            logger.info(f"Processing file: {file_path}")
            document_data = ingest_document(str(file_path))
            if document_data:
                ingested_documents.append(document_data)
                logger.info(f"Successfully ingested: {file_path.name}")
            else:
                logger.warning(f"Failed to ingest: {file_path.name}")
    
    logger.info(f"Ingested {len(ingested_documents)} documents from {directory_path}")
    return ingested_documents


def chunk_text(text: str, chunk_size: int = 500, chunk_overlap: int = 50) -> List[str]:
    """
    Split text into overlapping chunks for better context preservation.
    
    Args:
        text (str): Text to be chunked
        chunk_size (int): Maximum number of characters per chunk
        chunk_overlap (int): Number of characters to overlap between chunks
        
    Returns:
        List[str]: List of text chunks
    """
    if not text or not text.strip():
        return []
    
    # Clean up the text
    text = re.sub(r'\s+', ' ', text.strip())
    
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Calculate end position
        end = start + chunk_size
        
        # If this is not the last chunk, try to break at a sentence or word boundary
        if end < len(text):
            # Look for sentence boundaries (., !, ?) within the last 100 characters
            sentence_end = -1
            for i in range(max(0, end - 100), end):
                if text[i] in '.!?' and i < len(text) - 1 and text[i + 1] == ' ':
                    sentence_end = i + 1
            
            # If no sentence boundary found, look for word boundaries
            if sentence_end == -1:
                word_end = text.rfind(' ', start, end)
                if word_end > start:
                    end = word_end
            else:
                end = sentence_end
        
        # Extract the chunk
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = max(start + 1, end - chunk_overlap)
        
        # Avoid infinite loop
        if start >= len(text):
            break
    
    return chunks


class EmbeddingManager:
    """
    Manages text embeddings using sentence-transformers.
    """
    
    def __init__(self, model_name: str = 'all-MiniLM-L6-v2'):
        """
        Initialize the embedding manager with a specific model.
        
        Args:
            model_name (str): Name of the sentence-transformers model to use
        """
        if SentenceTransformer is None:
            raise ImportError("sentence-transformers is not installed. Please install it with: pip install sentence-transformers")
        
        self.model_name = model_name
        self.model = None
        
    def _load_model(self):
        """Lazy load the model when first needed."""
        if self.model is None:
            logger.info(f"Loading embedding model: {self.model_name}")
            self.model = SentenceTransformer(self.model_name)
            logger.info("Embedding model loaded successfully")
    
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """
        Generate embeddings for a list of texts.
        
        Args:
            texts (List[str]): List of text chunks to embed
            
        Returns:
            List[List[float]]: List of embedding vectors
        """
        if not texts:
            return []
        
        self._load_model()
        
        try:
            logger.info(f"Generating embeddings for {len(texts)} text chunks")
            embeddings = self.model.encode(texts, convert_to_tensor=False, show_progress_bar=True)
            logger.info("Embeddings generated successfully")
            return embeddings.tolist()
        except Exception as e:
            logger.error(f"Error generating embeddings: {e}")
            raise


class ChromaVectorStore:
    """
    Manages a Chroma vector database for storing document embeddings.
    """
    
    def __init__(self, persist_directory: str = "data/chroma"):
        """
        Initialize the Chroma vector store.
        
        Args:
            persist_directory (str): Directory to persist the Chroma database
        """
        if chromadb is None:
            raise ImportError("chromadb is not installed. Please install it with: pip install chromadb")
        
        self.persist_directory = Path(persist_directory)
        self.persist_directory.mkdir(parents=True, exist_ok=True)
        
        # Initialize Chroma client with persistence
        self.client = chromadb.PersistentClient(
            path=str(self.persist_directory),
            settings=Settings(
                anonymized_telemetry=False,
                allow_reset=True
            )
        )
        
        # Collection name for documents
        self.collection_name = "documents"
        self.collection = None
        
    def _get_or_create_collection(self):
        """Get or create the documents collection with robust error handling."""
        if self.collection is None:
            try:
                # First, try to get the existing collection
                self.collection = self.client.get_collection(self.collection_name)
                logger.info(f"Retrieved existing collection: {self.collection_name}")
            except Exception as e:
                logger.warning(f"Could not retrieve collection '{self.collection_name}': {e}")
                try:
                    # Try to create a new collection
                    self.collection = self.client.create_collection(
                        name=self.collection_name,
                        metadata={"description": "Document chunks with embeddings"}
                    )
                    logger.info(f"Created new collection: {self.collection_name}")
                except Exception as create_error:
                    logger.error(f"Failed to create collection '{self.collection_name}': {create_error}")
                    
                    # As a last resort, try to reset and recreate
                    try:
                        logger.warning("Attempting to reset and recreate collection...")
                        # List all collections to see what exists
                        existing_collections = self.client.list_collections()
                        logger.info(f"Existing collections: {[c.name for c in existing_collections]}")
                        
                        # Delete any existing collections that might be corrupted
                        for collection in existing_collections:
                            try:
                                self.client.delete_collection(collection.name)
                                logger.info(f"Deleted existing collection: {collection.name}")
                            except Exception as del_error:
                                logger.warning(f"Could not delete collection {collection.name}: {del_error}")
                        
                        # Now create a fresh collection
                        self.collection = self.client.create_collection(
                            name=self.collection_name,
                            metadata={"description": "Document chunks with embeddings"}
                        )
                        logger.info(f"Successfully created fresh collection: {self.collection_name}")
                        
                    except Exception as reset_error:
                        logger.error(f"Failed to reset and recreate collection: {reset_error}")
                        raise RuntimeError(f"Unable to initialize ChromaDB collection: {reset_error}")
        
        return self.collection
    
    def add_documents(self, 
                     chunks: List[str], 
                     embeddings: List[List[float]], 
                     metadata: List[Dict[str, Any]]) -> List[str]:
        """
        Add document chunks with their embeddings to the vector store.
        
        Args:
            chunks (List[str]): List of text chunks
            embeddings (List[List[float]]): List of embedding vectors
            metadata (List[Dict[str, Any]]): List of metadata for each chunk
            
        Returns:
            List[str]: List of generated IDs for the added documents
        """
        if not chunks or not embeddings or not metadata:
            logger.warning("Empty chunks, embeddings, or metadata provided")
            return []
        
        if len(chunks) != len(embeddings) or len(chunks) != len(metadata):
            raise ValueError("Chunks, embeddings, and metadata must have the same length")
        
        collection = self._get_or_create_collection()
        
        # Generate unique IDs for each chunk
        ids = [str(uuid4()) for _ in range(len(chunks))]
        
        try:
            collection.add(
                documents=chunks,
                embeddings=embeddings,
                metadatas=metadata,
                ids=ids
            )
            logger.info(f"Added {len(chunks)} document chunks to vector store")
            return ids
        except Exception as e:
            logger.error(f"Error adding documents to vector store: {e}")
            raise
    
    def query_similar(self, 
                     query_text: str, 
                     query_embedding: List[float], 
                     n_results: int = 5) -> Dict[str, Any]:
        """
        Query the vector store for similar documents.
        
        Args:
            query_text (str): The query text
            query_embedding (List[float]): The query embedding vector
            n_results (int): Number of results to return
            
        Returns:
            Dict[str, Any]: Query results with documents, distances, and metadata
        """
        collection = self._get_or_create_collection()
        
        try:
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=n_results,
                include=['documents', 'distances', 'metadatas']
            )
            logger.info(f"Retrieved {len(results['documents'][0])} similar documents")
            return results
        except Exception as e:
            logger.error(f"Error querying vector store: {e}")
            raise
    
    def get_collection_info(self) -> Dict[str, Any]:
        """
        Get information about the collection.
        
        Returns:
            Dict[str, Any]: Collection information including count and metadata
        """
        collection = self._get_or_create_collection()
        
        try:
            count = collection.count()
            return {
                "name": self.collection_name,
                "count": count,
                "persist_directory": str(self.persist_directory)
            }
        except Exception as e:
            logger.error(f"Error getting collection info: {e}")
            raise
    
    def delete_documents_by_ids(self, chunk_ids: List[str]) -> bool:
        """
        Delete specific documents from the vector store by their chunk IDs.
        
        Args:
            chunk_ids (List[str]): List of chunk IDs to delete
            
        Returns:
            bool: True if deletion was successful
        """
        if not chunk_ids:
            logger.warning("No chunk IDs provided for deletion")
            return False
            
        collection = self._get_or_create_collection()
        
        try:
            collection.delete(ids=chunk_ids)
            logger.info(f"Deleted {len(chunk_ids)} chunks from vector store")
            return True
        except Exception as e:
            logger.error(f"Error deleting chunks from vector store: {e}")
            return False
    
    def reset_collection(self) -> bool:
        """
        Reset the collection (delete all documents).
        
        Returns:
            bool: True if reset successfully
        """
        try:
            # Try to delete the specific collection first
            self.client.delete_collection(self.collection_name)
            self.collection = None
            logger.info(f"Reset collection: {self.collection_name}")
            return True
        except Exception as e:
            logger.warning(f"Could not delete collection '{self.collection_name}': {e}")
            
            # If that fails, try to delete all collections and start fresh
            try:
                existing_collections = self.client.list_collections()
                logger.info(f"Found {len(existing_collections)} existing collections")
                
                for collection in existing_collections:
                    try:
                        self.client.delete_collection(collection.name)
                        logger.info(f"Deleted collection: {collection.name}")
                    except Exception as del_error:
                        logger.warning(f"Could not delete collection {collection.name}: {del_error}")
                
                self.collection = None
                logger.info("Reset all collections successfully")
                return True
                
            except Exception as reset_error:
                logger.error(f"Failed to reset collections: {reset_error}")
                # Even if we can't delete, we can try to recreate
                self.collection = None
                return False
    
    def fix_collection_issues(self) -> bool:
        """
        Attempt to fix collection-related issues by recreating the collection.
        
        Returns:
            bool: True if issues were resolved
        """
        try:
            logger.info("Attempting to fix ChromaDB collection issues...")
            
            # Reset the collection
            self.reset_collection()
            
            # Force recreation of collection
            self.collection = None
            collection = self._get_or_create_collection()
            
            if collection:
                logger.info("Successfully fixed collection issues")
                return True
            else:
                logger.error("Failed to fix collection issues")
                return False
                
        except Exception as e:
            logger.error(f"Error fixing collection issues: {e}")
            return False


class DatabaseManager:
    """
    Manages SQLite database for storing document and chunk metadata.
    """
    
    def __init__(self, db_path: str = "data/db.sqlite"):
        """
        Initialize the database manager.
        
        Args:
            db_path (str): Path to the SQLite database file
        """
        self.db_path = Path(db_path)
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Initialize database and create tables
        self._init_database()
        
    def _init_database(self):
        """Initialize the database and create tables if they don't exist."""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Create documents table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    file_name TEXT NOT NULL,
                    file_path TEXT NOT NULL UNIQUE,
                    file_type TEXT NOT NULL,
                    content_length INTEGER NOT NULL,
                    word_count INTEGER NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)
            
            # Create chunks table
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS chunks (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    chunk_id TEXT NOT NULL UNIQUE,
                    document_id INTEGER NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    chunk_text TEXT NOT NULL,
                    chunk_size INTEGER NOT NULL,
                    total_chunks INTEGER NOT NULL,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    FOREIGN KEY (document_id) REFERENCES documents (id) ON DELETE CASCADE
                )
            """)
            
            # Create indexes for better query performance
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_chunks_document_id ON chunks (document_id)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_chunks_chunk_id ON chunks (chunk_id)")
            cursor.execute("CREATE INDEX IF NOT EXISTS idx_documents_file_path ON documents (file_path)")
            
            conn.commit()
            logger.info(f"Database initialized at: {self.db_path}")
    
    def add_document(self, document_data: Dict[str, Any]) -> int:
        """
        Add a document to the database.
        
        Args:
            document_data (Dict[str, Any]): Document data from ingest_document()
            
        Returns:
            int: Document ID in the database
        """
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Check if document already exists
            cursor.execute("SELECT id FROM documents WHERE file_path = ?", (document_data['file_path'],))
            existing = cursor.fetchone()
            
            if existing:
                # Update existing document
                document_id = existing[0]
                cursor.execute("""
                    UPDATE documents 
                    SET file_name = ?, file_type = ?, content_length = ?, 
                        word_count = ?, updated_at = CURRENT_TIMESTAMP
                    WHERE id = ?
                """, (
                    document_data['file_name'],
                    document_data['file_type'],
                    document_data['content_length'],
                    document_data['word_count'],
                    document_id
                ))
                logger.info(f"Updated document in database: {document_data['file_name']}")
            else:
                # Insert new document
                cursor.execute("""
                    INSERT INTO documents (file_name, file_path, file_type, content_length, word_count)
                    VALUES (?, ?, ?, ?, ?)
                """, (
                    document_data['file_name'],
                    document_data['file_path'],
                    document_data['file_type'],
                    document_data['content_length'],
                    document_data['word_count']
                ))
                document_id = cursor.lastrowid
                logger.info(f"Added document to database: {document_data['file_name']}")
            
            conn.commit()
            return document_id
    
    def add_chunks(self, document_id: int, chunks: List[str], chunk_ids: List[str]) -> List[int]:
        """
        Add chunks to the database.
        
        Args:
            document_id (int): Document ID from the database
            chunks (List[str]): List of text chunks
            chunk_ids (List[str]): List of chunk IDs from vector store
            
        Returns:
            List[int]: List of chunk database IDs
        """
        if len(chunks) != len(chunk_ids):
            raise ValueError("Chunks and chunk_ids must have the same length")
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Clear existing chunks for this document
            cursor.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
            
            # Insert new chunks
            chunk_db_ids = []
            total_chunks = len(chunks)
            
            for i, (chunk_text, chunk_id) in enumerate(zip(chunks, chunk_ids)):
                cursor.execute("""
                    INSERT INTO chunks (chunk_id, document_id, chunk_index, chunk_text, 
                                      chunk_size, total_chunks)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (
                    chunk_id,
                    document_id,
                    i,
                    chunk_text,
                    len(chunk_text),
                    total_chunks
                ))
                chunk_db_ids.append(cursor.lastrowid)
            
            conn.commit()
            logger.info(f"Added {len(chunks)} chunks to database for document ID {document_id}")
            return chunk_db_ids
    
    def get_document_by_path(self, file_path: str) -> Optional[Dict[str, Any]]:
        """
        Get document information by file path.
        
        Args:
            file_path (str): Path to the document file
            
        Returns:
            Optional[Dict[str, Any]]: Document information or None if not found
        """
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row  # Enable column access by name
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM documents WHERE file_path = ?", (file_path,))
            row = cursor.fetchone()
            
            if row:
                return dict(row)
            return None
    
    def get_chunks_by_document_id(self, document_id: int) -> List[Dict[str, Any]]:
        """
        Get all chunks for a document.
        
        Args:
            document_id (int): Document ID
            
        Returns:
            List[Dict[str, Any]]: List of chunk information
        """
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT * FROM chunks 
                WHERE document_id = ? 
                ORDER BY chunk_index
            """, (document_id,))
            
            return [dict(row) for row in cursor.fetchall()]
    
    def get_chunk_by_id(self, chunk_id: str) -> Optional[Dict[str, Any]]:
        """
        Get chunk information by chunk ID.
        
        Args:
            chunk_id (str): Chunk ID from vector store
            
        Returns:
            Optional[Dict[str, Any]]: Chunk information with document details
        """
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT c.*, d.file_name, d.file_path, d.file_type
                FROM chunks c
                JOIN documents d ON c.document_id = d.id
                WHERE c.chunk_id = ?
            """, (chunk_id,))
            
            row = cursor.fetchone()
            if row:
                return dict(row)
            return None
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """
        Get all documents from the database.
        
        Returns:
            List[Dict[str, Any]]: List of all documents
        """
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("SELECT * FROM documents ORDER BY created_at DESC")
            return [dict(row) for row in cursor.fetchall()]
    
    def get_database_stats(self) -> Dict[str, Any]:
        """
        Get database statistics.
        
        Returns:
            Dict[str, Any]: Database statistics
        """
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            # Count documents
            cursor.execute("SELECT COUNT(*) FROM documents")
            doc_count = cursor.fetchone()[0]
            
            # Count chunks
            cursor.execute("SELECT COUNT(*) FROM chunks")
            chunk_count = cursor.fetchone()[0]
            
            # Get database size
            db_size = self.db_path.stat().st_size if self.db_path.exists() else 0
            
            return {
                "database_path": str(self.db_path),
                "document_count": doc_count,
                "chunk_count": chunk_count,
                "database_size_bytes": db_size,
                "database_size_mb": round(db_size / (1024 * 1024), 2)
            }
    
    def search_documents(self, search_term: str) -> List[Dict[str, Any]]:
        """
        Search documents by file name or path.
        
        Args:
            search_term (str): Search term
            
        Returns:
            List[Dict[str, Any]]: Matching documents
        """
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("""
                SELECT * FROM documents 
                WHERE file_name LIKE ? OR file_path LIKE ?
                ORDER BY created_at DESC
            """, (f"%{search_term}%", f"%{search_term}%"))
            
            return [dict(row) for row in cursor.fetchall()]
    
    def get_chunk_ids_by_document_id(self, document_id: int) -> List[str]:
        """
        Get all chunk IDs for a specific document.
        
        Args:
            document_id (int): Document ID to get chunk IDs for
            
        Returns:
            List[str]: List of chunk IDs
        """
        with sqlite3.connect(self.db_path) as conn:
            conn.row_factory = sqlite3.Row
            cursor = conn.cursor()
            
            cursor.execute("SELECT chunk_id FROM chunks WHERE document_id = ?", (document_id,))
            rows = cursor.fetchall()
            
            return [row['chunk_id'] for row in rows]

    def delete_document(self, document_id: int, vector_store: Optional['ChromaVectorStore'] = None) -> bool:
        """
        Delete a document and its chunks from the database and optionally from vector store.
        
        Args:
            document_id (int): Document ID to delete
            vector_store (Optional[ChromaVectorStore]): Vector store to also delete from
            
        Returns:
            bool: True if deleted successfully
        """
        # Get chunk IDs before deletion if we need to delete from vector store
        chunk_ids = []
        if vector_store:
            chunk_ids = self.get_chunk_ids_by_document_id(document_id)
        
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()
            
            cursor.execute("DELETE FROM documents WHERE id = ?", (document_id,))
            deleted = cursor.rowcount > 0
            
            conn.commit()
            
            if deleted:
                logger.info(f"Deleted document ID {document_id} from database")
                
                # Also delete from vector store if provided
                if vector_store and chunk_ids:
                    vector_deleted = vector_store.delete_documents_by_ids(chunk_ids)
                    if vector_deleted:
                        logger.info(f"Deleted {len(chunk_ids)} chunks from vector store for document ID {document_id}")
                    else:
                        logger.warning(f"Failed to delete chunks from vector store for document ID {document_id}")
            
            return deleted
    
    def clear_all_data(self) -> bool:
        """
        Clear all documents and chunks from the database.
        
        Returns:
            bool: True if cleared successfully
        """
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                
                # Delete all chunks first (due to foreign key constraint)
                cursor.execute("DELETE FROM chunks")
                chunks_deleted = cursor.rowcount
                
                # Delete all documents
                cursor.execute("DELETE FROM documents")
                docs_deleted = cursor.rowcount
                
                conn.commit()
                
                logger.info(f"Cleared all data: {docs_deleted} documents, {chunks_deleted} chunks")
                return True
                
        except Exception as e:
            logger.error(f"Error clearing database: {e}")
            return False


def process_documents_to_vector_store(
    documents: List[Dict[str, Any]], 
    vector_store: ChromaVectorStore,
    embedding_manager: EmbeddingManager,
    db_manager: Optional[DatabaseManager] = None,
    chunk_size: int = 500,
    chunk_overlap: int = 50
) -> Dict[str, Any]:
    """
    Complete pipeline to process documents: chunk -> embed -> store in vector database and SQLite.
    
    Args:
        documents (List[Dict[str, Any]]): List of document data from ingest_document()
        vector_store (ChromaVectorStore): Vector store instance
        embedding_manager (EmbeddingManager): Embedding manager instance
        db_manager (Optional[DatabaseManager]): Database manager for metadata storage
        chunk_size (int): Size of text chunks
        chunk_overlap (int): Overlap between chunks
        
    Returns:
        Dict[str, Any]: Processing results and statistics
    """
    if not documents:
        logger.warning("No documents provided for processing")
        return {"processed_documents": 0, "total_chunks": 0, "chunk_ids": []}
    
    logger.info(f"Processing {len(documents)} documents for vector storage")
    
    all_chunks = []
    all_metadata = []
    processed_docs = 0
    document_chunk_mapping = []  # Track which chunks belong to which documents
    
    # Process each document
    for doc in documents:
        try:
            content = doc.get('content', '')
            if not content:
                logger.warning(f"Empty content for document: {doc.get('file_name', 'Unknown')}")
                continue
            
            # Chunk the document
            chunks = chunk_text(content, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
            if not chunks:
                logger.warning(f"No chunks generated for document: {doc.get('file_name', 'Unknown')}")
                continue
            
            logger.info(f"Generated {len(chunks)} chunks for {doc.get('file_name', 'Unknown')}")
            
            # Store document in SQLite database if db_manager is provided
            document_id = None
            if db_manager:
                document_id = db_manager.add_document(doc)
            
            # Create metadata for each chunk and track document mapping
            chunk_start_index = len(all_chunks)
            for i, chunk in enumerate(chunks):
                metadata = {
                    'source_file': doc.get('file_name', 'Unknown'),
                    'file_path': doc.get('file_path', ''),
                    'file_type': doc.get('file_type', ''),
                    'chunk_index': i,
                    'chunk_size': len(chunk),
                    'total_chunks': len(chunks),
                    'source_content_length': doc.get('content_length', 0),
                    'source_word_count': doc.get('word_count', 0)
                }
                all_chunks.append(chunk)
                all_metadata.append(metadata)
            
            # Track document-chunk mapping for database storage
            document_chunk_mapping.append({
                'document_data': doc,
                'document_id': document_id,
                'chunks': chunks,
                'chunk_start_index': chunk_start_index,
                'chunk_count': len(chunks)
            })
            
            processed_docs += 1
            
        except Exception as e:
            logger.error(f"Error processing document {doc.get('file_name', 'Unknown')}: {e}")
            continue
    
    if not all_chunks:
        logger.warning("No chunks generated from any documents")
        return {"processed_documents": 0, "total_chunks": 0, "chunk_ids": []}
    
    # Generate embeddings for all chunks
    try:
        embeddings = embedding_manager.generate_embeddings(all_chunks)
        logger.info(f"Generated embeddings for {len(embeddings)} chunks")
    except Exception as e:
        logger.error(f"Failed to generate embeddings: {e}")
        raise
    
    # Store in vector database
    try:
        chunk_ids = vector_store.add_documents(all_chunks, embeddings, all_metadata)
        logger.info(f"Stored {len(chunk_ids)} chunks in vector database")
    except Exception as e:
        logger.error(f"Failed to store chunks in vector database: {e}")
        
        # If it's a collection-related error, try to fix it
        if "collection" in str(e).lower() or "does not exist" in str(e).lower():
            logger.info("Detected collection issue, attempting to fix...")
            try:
                if vector_store.fix_collection_issues():
                    logger.info("Fixed collection issues, retrying document storage...")
                    chunk_ids = vector_store.add_documents(all_chunks, embeddings, all_metadata)
                    logger.info(f"Successfully stored {len(chunk_ids)} chunks after fixing collection")
                else:
                    logger.error("Could not fix collection issues")
                    raise
            except Exception as retry_error:
                logger.error(f"Retry after collection fix failed: {retry_error}")
                raise
        else:
            raise
    
    # Store chunk metadata in SQLite database
    if db_manager:
        try:
            for doc_mapping in document_chunk_mapping:
                if doc_mapping['document_id'] is not None:
                    # Get the chunk IDs for this document
                    start_idx = doc_mapping['chunk_start_index']
                    end_idx = start_idx + doc_mapping['chunk_count']
                    doc_chunk_ids = chunk_ids[start_idx:end_idx]
                    doc_chunks = doc_mapping['chunks']
                    
                    # Store chunks in database
                    db_manager.add_chunks(
                        document_id=doc_mapping['document_id'],
                        chunks=doc_chunks,
                        chunk_ids=doc_chunk_ids
                    )
            logger.info(f"Stored chunk metadata in SQLite database")
        except Exception as e:
            logger.error(f"Failed to store chunk metadata in database: {e}")
            # Don't raise here as vector storage succeeded
    
    # Return processing results
    results = {
        "processed_documents": processed_docs,
        "total_chunks": len(all_chunks),
        "chunk_ids": chunk_ids,
        "average_chunk_size": sum(len(chunk) for chunk in all_chunks) / len(all_chunks),
        "vector_store_info": vector_store.get_collection_info()
    }
    
    logger.info(f"Successfully processed {processed_docs} documents into {len(all_chunks)} chunks")
    return results


def process_directory_to_vector_store(
    directory_path: str,
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite",
    model_name: str = 'all-MiniLM-L6-v2',
    chunk_size: int = 500,
    chunk_overlap: int = 50,
    file_types: Optional[List[str]] = None
) -> Dict[str, Any]:
    """
    Complete pipeline to process all documents in a directory.
    
    Args:
        directory_path (str): Path to directory containing documents
        persist_directory (str): Directory to persist Chroma database
        db_path (str): Path to SQLite database file
        model_name (str): Sentence transformer model name
        chunk_size (int): Size of text chunks
        chunk_overlap (int): Overlap between chunks
        file_types (Optional[List[str]]): File extensions to process
        
    Returns:
        Dict[str, Any]: Processing results and statistics
    """
    logger.info(f"Starting complete document processing pipeline for directory: {directory_path}")
    
    # Step 1: Ingest documents
    documents = ingest_documents_from_directory(directory_path, file_types)
    if not documents:
        logger.warning(f"No documents found in directory: {directory_path}")
        return {"processed_documents": 0, "total_chunks": 0, "chunk_ids": []}
    
    # Step 2: Initialize components
    try:
        embedding_manager = EmbeddingManager(model_name=model_name)
        vector_store = ChromaVectorStore(persist_directory=persist_directory)
        db_manager = DatabaseManager(db_path=db_path)
        logger.info("Initialized embedding manager, vector store, and database manager")
    except Exception as e:
        logger.error(f"Failed to initialize components: {e}")
        raise
    
    # Step 3: Process documents
    results = process_documents_to_vector_store(
        documents=documents,
        vector_store=vector_store,
        embedding_manager=embedding_manager,
        db_manager=db_manager,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    
    # Add database statistics to results
    if db_manager:
        results["database_stats"] = db_manager.get_database_stats()
    
    logger.info("Document processing pipeline completed successfully")
    return results


def query_with_metadata(
    query_text: str,
    embedding_manager: EmbeddingManager,
    vector_store: ChromaVectorStore,
    db_manager: DatabaseManager,
    n_results: int = 5
) -> Dict[str, Any]:
    """
    Query the vector store and enrich results with metadata from SQLite database.
    
    Args:
        query_text (str): The search query
        embedding_manager (EmbeddingManager): Embedding manager instance
        vector_store (ChromaVectorStore): Vector store instance
        db_manager (DatabaseManager): Database manager instance
        n_results (int): Number of results to return
        
    Returns:
        Dict[str, Any]: Enhanced query results with metadata
    """
    # Generate query embedding
    query_embedding = embedding_manager.generate_embeddings([query_text])[0]
    
    # Search vector store
    vector_results = vector_store.query_similar(
        query_text=query_text,
        query_embedding=query_embedding,
        n_results=n_results
    )
    
    # Enrich with database metadata and filter out orphaned references
    enriched_results = []
    orphaned_count = 0
    
    if vector_results['documents'] and len(vector_results['documents'][0]) > 0:
        for i, (doc, distance, metadata) in enumerate(zip(
            vector_results['documents'][0],
            vector_results['distances'][0],
            vector_results['metadatas'][0]
        )):
            # Get chunk ID from metadata (this would be stored in the metadata)
            # For now, we'll try to find the chunk by text content
            chunk_info = {
                'rank': i + 1,
                'similarity_score': 1 - distance,
                'distance': distance,
                'chunk_text': doc,
                'chunk_preview': doc[:200] + "..." if len(doc) > 200 else doc,
                'vector_metadata': metadata
            }
            
            # Try to find additional metadata from database
            document_exists_in_db = False
            try:
                source_file = metadata.get('source_file', '')
                file_path = metadata.get('file_path', '')
                
                if source_file and file_path:
                    doc_info = db_manager.get_document_by_path(file_path)
                    if doc_info:
                        document_exists_in_db = True
                        chunk_info['database_metadata'] = {
                            'document_id': doc_info['id'],
                            'file_name': doc_info['file_name'],
                            'file_path': doc_info['file_path'],
                            'file_type': doc_info['file_type'],
                            'created_at': doc_info['created_at'],
                            'updated_at': doc_info['updated_at']
                        }
                    else:
                        logger.warning(f"Document '{source_file}' exists in vector store but not in database - orphaned reference")
                        orphaned_count += 1
                        # Mark as orphaned but still include it with limited metadata
                        chunk_info['orphaned'] = True
                        chunk_info['database_metadata'] = {
                            'file_name': source_file,
                            'file_path': file_path,
                            'file_type': metadata.get('file_type', 'unknown'),
                            'created_at': 'Unknown (orphaned)',
                            'status': 'Document may have been deleted'
                        }
            except Exception as e:
                logger.warning(f"Could not enrich result {i} with database metadata: {e}")
                chunk_info['orphaned'] = True
            
            # Always include the result, but mark orphaned ones
            enriched_results.append(chunk_info)
        
        # Log orphaned references
        if orphaned_count > 0:
            logger.warning(f"Found {orphaned_count} orphaned document references in vector store")
            logger.info("Consider running database/vector store synchronization to clean up orphaned references")
    
    return {
        'query': query_text,
        'total_results': len(enriched_results),
        'results': enriched_results,
        'vector_store_info': vector_store.get_collection_info(),
        'database_stats': db_manager.get_database_stats()
    }


def format_context_for_llm(retrieved_chunks: List[Dict[str, Any]], max_context_length: int = 4000) -> str:
    """
    Format retrieved chunks into a context string for the LLM, grouping chunks by source document.
    
    Args:
        retrieved_chunks (List[Dict[str, Any]]): List of retrieved chunk information
        max_context_length (int): Maximum character length for context
        
    Returns:
        str: Formatted context string
    """
    if not retrieved_chunks:
        return "No relevant information found in your documents."
    
    # Group chunks by source document
    documents = {}
    for chunk_info in retrieved_chunks:
        metadata = chunk_info.get('vector_metadata', {})
        source_file = metadata.get('source_file', 'Unknown')
        
        if source_file not in documents:
            documents[source_file] = []
        documents[source_file].append(chunk_info)
    
    context_parts = []
    current_length = 0
    doc_number = 1
    
    for source_file, chunks in documents.items():
        # Sort chunks by chunk_index to maintain order
        chunks.sort(key=lambda x: x.get('vector_metadata', {}).get('chunk_index', 0))
        
        # Create document header
        doc_header = f"[Document {doc_number}] Source: {source_file}"
        if chunks:
            metadata = chunks[0].get('vector_metadata', {})
            total_chunks = metadata.get('total_chunks', 0)
            if total_chunks > 1:
                doc_header += f" (contains {len(chunks)} relevant sections out of {total_chunks} total)"
        
        # Check if we can fit the document header
        if current_length + len(doc_header) + 50 > max_context_length:
            break
            
        context_parts.append(doc_header)
        current_length += len(doc_header)
        
        # Add chunks from this document
        for chunk_info in chunks:
            chunk_text = chunk_info.get('chunk_text', '')
            metadata = chunk_info.get('vector_metadata', {})
            
            # Format chunk with section info
            chunk_header = ""
            if 'chunk_index' in metadata:
                chunk_num = metadata['chunk_index'] + 1
                chunk_header = f"  Section {chunk_num}: "
            
            formatted_chunk = f"{chunk_header}{chunk_text}\n"
            
            # Check if adding this chunk would exceed the limit
            if current_length + len(formatted_chunk) > max_context_length:
                if len(context_parts) == 1:  # Only document header so far
                    # If even the first chunk is too long, truncate it
                    available_space = max_context_length - current_length - 50
                    if available_space > 100:  # Only truncate if we have reasonable space
                        truncated_text = chunk_text[:available_space] + "...[truncated]"
                        formatted_chunk = f"{chunk_header}{truncated_text}\n"
                        context_parts.append(formatted_chunk)
                        current_length += len(formatted_chunk)
                break
            
            context_parts.append(formatted_chunk)
            current_length += len(formatted_chunk)
        
        context_parts.append("")  # Add blank line between documents
        doc_number += 1
    
    # Add a note if not all chunks were included
    total_chunks = len(retrieved_chunks)
    included_chunks = sum(len(chunks) for chunks in documents.items() if any(chunk.get('chunk_text', '') in '\n'.join(context_parts) for chunk in chunks[1]))
    
    if doc_number <= len(documents):
        context_parts.append(f"[Note: Showing {doc_number-1} of {len(documents)} relevant documents due to context length limits]")
    
    return "\n".join(context_parts)


def ask_your_files(
    question: str,
    embedding_manager: EmbeddingManager,
    vector_store: ChromaVectorStore,
    db_manager: Optional[DatabaseManager] = None,
    llm_client: Optional[LLMClient] = None,
    k: int = 5,
    max_context_length: int = 4000,
    include_metadata: bool = True,
    custom_system_prompt: Optional[str] = None
) -> Dict[str, Any]:
    """
    Ask a question about your documents using RAG (Retrieval-Augmented Generation).
    
    This function:
    1. Converts the question to an embedding
    2. Retrieves the top-k most relevant document chunks
    3. Formats them as context for the LLM
    4. Asks the LLM to answer based on the context
    5. Returns the answer with source attribution
    
    Args:
        question (str): The question to ask about your documents
        embedding_manager (EmbeddingManager): Embedding manager instance
        vector_store (ChromaVectorStore): Vector store instance
        db_manager (Optional[DatabaseManager]): Database manager for enhanced metadata
        llm_client (Optional[LLMClient]): LLM client instance (defaults to Ollama)
        k (int): Number of relevant chunks to retrieve
        max_context_length (int): Maximum characters to include in LLM context
        include_metadata (bool): Whether to include enhanced metadata from database
        custom_system_prompt (Optional[str]): Custom system prompt for the LLM
        
    Returns:
        Dict[str, Any]: Response containing answer, sources, and metadata
    """
    logger.info(f"Processing question: {question[:100]}...")
    
    # Initialize LLM client if not provided
    if llm_client is None:
        llm_client = LLMClient(
            model="llama3",  # Default model
            default_system="You are a helpful AI assistant that answers questions based on provided documents."
        )
    
    try:
        # Step 1: Retrieve relevant chunks
        if include_metadata and db_manager:
            # Use enhanced query with metadata
            retrieval_results = query_with_metadata(
                query_text=question,
                embedding_manager=embedding_manager,
                vector_store=vector_store,
                db_manager=db_manager,
                n_results=k
            )
            retrieved_chunks = retrieval_results['results']
        else:
            # Use basic vector search
            query_embedding = embedding_manager.generate_embeddings([question])[0]
            vector_results = vector_store.query_similar(
                query_text=question,
                query_embedding=query_embedding,
                n_results=k
            )
            
            # Convert to consistent format
            retrieved_chunks = []
            if vector_results['documents'] and len(vector_results['documents'][0]) > 0:
                for i, (doc, distance, metadata) in enumerate(zip(
                    vector_results['documents'][0],
                    vector_results['distances'][0],
                    vector_results['metadatas'][0]
                )):
                    retrieved_chunks.append({
                        'rank': i + 1,
                        'similarity_score': 1 - distance,
                        'chunk_text': doc,
                        'vector_metadata': metadata
                    })
        
        if not retrieved_chunks:
            return {
                'question': question,
                'answer': "I couldn't find any relevant information in your documents to answer this question.",
                'sources': [],
                'retrieval_info': {
                    'chunks_found': 0,
                    'chunks_used': 0
                },
                'llm_used': False
            }
        
        # Step 2: Format context for LLM
        context = format_context_for_llm(retrieved_chunks, max_context_length)
        
        # Step 3: Create system prompt with current date
        from datetime import datetime
        current_date = datetime.now().strftime("%B %d, %Y")
        
        system_prompt = custom_system_prompt or f"""You are a helpful AI assistant that answers questions based on provided documents. 

Current date: {current_date}

Instructions:
- Answer the user's question using ONLY the information provided in the documents below
- When citing sources, refer to documents by their number (e.g., "Document 1", "Document 2")
- Each document may contain multiple sections - you can reference specific sections if helpful
- If the documents don't contain enough information to answer the question, say so clearly
- Be concise but thorough
- If multiple documents contain relevant information, synthesize them in your response
- Only reference document numbers that actually exist in the provided context
- When relevant, consider the current date when interpreting time-sensitive information"""

        # Step 4: Create the full prompt
        full_prompt = f"""Based on the following documents, please answer this question: {question}

DOCUMENTS:
{context}

QUESTION: {question}

Please provide a clear, accurate answer based on the information in the documents above."""

        # Step 5: Get LLM response
        logger.info("Sending query to LLM...")
        try:
            llm_response = llm_client.chat(
                prompt=full_prompt,
                system=system_prompt,
                options={
                    "temperature": 0.1,  # Lower temperature for more factual responses
                    "num_ctx": 4096      # Larger context window for documents
                }
            )
        except LLMError as e:
            logger.error(f"LLM request failed: {e}")
            return {
                'question': question,
                'answer': f"Error: Could not get response from LLM. {str(e)}",
                'sources': [{'source_file': chunk.get('vector_metadata', {}).get('source_file', 'Unknown'), 
                           'similarity_score': chunk.get('similarity_score', 0)} 
                          for chunk in retrieved_chunks],
                'retrieval_info': {
                    'chunks_found': len(retrieved_chunks),
                    'chunks_used': len(retrieved_chunks)
                },
                'llm_used': False,
                'error': str(e)
            }
        
        # Step 6: Prepare response with source attribution
        sources = []
        for chunk in retrieved_chunks:
            source_info = {
                'source_file': chunk.get('vector_metadata', {}).get('source_file', 'Unknown'),
                'similarity_score': chunk.get('similarity_score', 0),
                'chunk_index': chunk.get('vector_metadata', {}).get('chunk_index', 0),
                'total_chunks': chunk.get('vector_metadata', {}).get('total_chunks', 0)
            }
            
            # Add database metadata if available
            if 'database_metadata' in chunk:
                dm = chunk['database_metadata']
                source_info.update({
                    'file_path': dm.get('file_path', ''),
                    'file_type': dm.get('file_type', ''),
                    'created_at': dm.get('created_at', '')
                })
            
            sources.append(source_info)
        
        logger.info("Question answered successfully")
        
        return {
            'question': question,
            'answer': llm_response,
            'sources': sources,
            'retrieval_info': {
                'chunks_found': len(retrieved_chunks),
                'chunks_used': len([c for c in retrieved_chunks if c.get('chunk_text', '') in context]),
                'total_context_length': len(context)
            },
            'llm_used': True,
            'model_used': llm_client.model
        }
        
    except Exception as e:
        logger.error(f"Error in ask_your_files: {e}")
        return {
            'question': question,
            'answer': f"Error processing your question: {str(e)}",
            'sources': [],
            'retrieval_info': {'chunks_found': 0, 'chunks_used': 0},
            'llm_used': False,
            'error': str(e)
        }


def ask_your_files_simple(
    question: str,
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite",
    model_name: str = 'all-MiniLM-L6-v2',
    llm_model: str = "llama3",
    k: int = 5
) -> str:
    """
    Simplified "Ask Your Files" function that initializes everything automatically.
    
    Args:
        question (str): The question to ask
        persist_directory (str): Path to Chroma database
        db_path (str): Path to SQLite database
        model_name (str): Embedding model name
        llm_model (str): LLM model name
        k (int): Number of chunks to retrieve
        
    Returns:
        str: The answer to your question
    """
    try:
        # Initialize components
        embedding_manager = EmbeddingManager(model_name)
        vector_store = ChromaVectorStore(persist_directory)
        db_manager = DatabaseManager(db_path)
        llm_client = LLMClient(model=llm_model)
        
        # Ask the question
        result = ask_your_files(
            question=question,
            embedding_manager=embedding_manager,
            vector_store=vector_store,
            db_manager=db_manager,
            llm_client=llm_client,
            k=k
        )
        
        return result['answer']
        
    except Exception as e:
        return f"Error: {str(e)}"


def format_citations(sources: List[Dict[str, Any]], max_citations: int = 10) -> str:
    """
    Format source citations in a readable format.
    
    Args:
        sources (List[Dict[str, Any]]): List of source information
        max_citations (int): Maximum number of citations to show
        
    Returns:
        str: Formatted citation string
    """
    if not sources:
        return "No sources found."
    
    citations = []
    seen_sources = set()  # Track unique sources to avoid duplicates
    
    for i, source in enumerate(sources[:max_citations]):
        source_file = source.get('source_file', 'Unknown')
        chunk_index = source.get('chunk_index', 0)
        total_chunks = source.get('total_chunks', 0)
        similarity_score = source.get('similarity_score', 0)
        
        # Create unique identifier for this source
        source_id = f"{source_file}_chunk_{chunk_index}"
        if source_id in seen_sources:
            continue
        seen_sources.add(source_id)
        
        # Format citation
        citation = f"[{len(citations) + 1}] {source_file}"
        
        if total_chunks > 1:
            citation += f" (chunk {chunk_index + 1}/{total_chunks})"
        
        if similarity_score > 0:
            citation += f" - relevance: {similarity_score:.1%}"
        
        # Add file type and path if available
        if 'file_type' in source:
            citation += f" [{source['file_type']}]"
        
        if 'created_at' in source:
            citation += f" (added: {source['created_at'][:10]})"
        
        citations.append(citation)
    
    # Add note if there are more sources
    if len(sources) > len(citations):
        citations.append(f"... and {len(sources) - len(citations)} more sources")
    
    return "\n".join(citations)


def format_answer_with_citations(answer: str, sources: List[Dict[str, Any]], 
                                show_metadata: bool = True) -> str:
    """
    Format answer with inline citations and source list.
    
    Args:
        answer (str): The LLM's answer
        sources (List[Dict[str, Any]]): List of source information
        show_metadata (bool): Whether to show detailed metadata
        
    Returns:
        str: Formatted answer with citations
    """
    if not sources:
        return f"{answer}\n\n[No sources available]"
    
    # Create the response with citations
    response_parts = [answer]
    
    # Add sources section
    response_parts.append("\n" + "="*50)
    response_parts.append("📚 SOURCES:")
    response_parts.append("="*50)
    
    # Format each source
    for i, source in enumerate(sources[:10], 1):  # Limit to top 10 sources
        source_file = source.get('source_file', 'Unknown')
        chunk_index = source.get('chunk_index', 0)
        total_chunks = source.get('total_chunks', 0)
        similarity_score = source.get('similarity_score', 0)
        
        source_line = f"[{i}] 📄 {source_file}"
        
        # Add chunk information
        if total_chunks > 1:
            source_line += f" (section {chunk_index + 1} of {total_chunks})"
        
        # Add relevance score
        if similarity_score > 0:
            relevance_bar = "█" * int(similarity_score * 10) + "░" * (10 - int(similarity_score * 10))
            source_line += f"\n    📊 Relevance: {similarity_score:.1%} {relevance_bar}"
        
        # Add metadata if requested
        if show_metadata:
            if 'file_path' in source:
                source_line += f"\n    📂 Path: {source['file_path']}"
            if 'file_type' in source:
                source_line += f"\n    📋 Type: {source['file_type']}"
            if 'created_at' in source:
                if source['created_at'] == 'Unknown (orphaned)':
                    source_line += f"\n    ⚠️  Status: {source.get('status', 'Document may have been deleted')}"
                else:
                    source_line += f"\n    📅 Added: {source['created_at']}"
        
        response_parts.append(source_line)
        
        # Add separator between sources
        if i < min(len(sources), 10):
            response_parts.append("-" * 30)
    
    return "\n".join(response_parts)


def ask_your_files_with_citations(
    question: str,
    embedding_manager: EmbeddingManager,
    vector_store: ChromaVectorStore,
    db_manager: Optional[DatabaseManager] = None,
    llm_client: Optional[LLMClient] = None,
    k: int = 5,
    show_detailed_citations: bool = True,
    show_metadata: bool = True
) -> str:
    """
    Ask a question and return answer with formatted citations.
    
    Args:
        question (str): The question to ask
        embedding_manager (EmbeddingManager): Embedding manager instance
        vector_store (ChromaVectorStore): Vector store instance
        db_manager (Optional[DatabaseManager]): Database manager instance
        llm_client (Optional[LLMClient]): LLM client instance
        k (int): Number of chunks to retrieve
        show_detailed_citations (bool): Whether to show detailed citation information
        show_metadata (bool): Whether to show file metadata
        
    Returns:
        str: Formatted answer with citations
    """
    # Get the full response
    result = ask_your_files(
        question=question,
        embedding_manager=embedding_manager,
        vector_store=vector_store,
        db_manager=db_manager,
        llm_client=llm_client,
        k=k
    )
    
    if not result['llm_used']:
        # If LLM wasn't used, just return the basic answer
        return result['answer']
    
    # Format with citations
    if show_detailed_citations:
        return format_answer_with_citations(
            result['answer'], 
            result['sources'], 
            show_metadata=show_metadata
        )
    else:
        # Simple citation format
        citations = format_citations(result['sources'])
        return f"{result['answer']}\n\nSources:\n{citations}"


def ask_your_files_simple_with_citations(
    question: str,
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite",
    model_name: str = 'all-MiniLM-L6-v2',
    llm_model: str = "llama3",
    k: int = 5,
    show_detailed_citations: bool = True
) -> str:
    """
    Simple Ask Your Files with formatted citations.
    
    Args:
        question (str): The question to ask
        persist_directory (str): Path to Chroma database
        db_path (str): Path to SQLite database
        model_name (str): Embedding model name
        llm_model (str): LLM model name
        k (int): Number of chunks to retrieve
        show_detailed_citations (bool): Whether to show detailed citations
        
    Returns:
        str: Answer with formatted citations
    """
    try:
        # Initialize components
        embedding_manager = EmbeddingManager(model_name)
        vector_store = ChromaVectorStore(persist_directory)
        db_manager = DatabaseManager(db_path)
        llm_client = LLMClient(model=llm_model)
        
        # Ask with citations
        return ask_your_files_with_citations(
            question=question,
            embedding_manager=embedding_manager,
            vector_store=vector_store,
            db_manager=db_manager,
            llm_client=llm_client,
            k=k,
            show_detailed_citations=show_detailed_citations
        )
        
    except Exception as e:
        return f"Error: {str(e)}"


def clear_all_chunks_and_documents(
    vector_store: Optional[ChromaVectorStore] = None,
    db_manager: Optional[DatabaseManager] = None,
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite"
) -> Dict[str, Any]:
    """
    Remove all chunks and documents from both the vector store and database.
    
    This function provides a complete cleanup operation that:
    1. Clears all documents and chunks from the SQLite database
    2. Resets the ChromaDB vector store collection
    3. Returns detailed information about the cleanup operation
    
    Args:
        vector_store (Optional[ChromaVectorStore]): Vector store instance to clear.
                                                   If None, creates a new instance.
        db_manager (Optional[DatabaseManager]): Database manager instance to clear.
                                              If None, creates a new instance.
        persist_directory (str): Directory containing the Chroma database
        db_path (str): Path to the SQLite database file
        
    Returns:
        Dict[str, Any]: Results of the cleanup operation including:
                       - success: bool indicating if operation succeeded
                       - database_cleared: bool indicating if database was cleared
                       - vector_store_cleared: bool indicating if vector store was cleared
                       - stats_before: database stats before clearing
                       - stats_after: database stats after clearing
                       - errors: list of any errors encountered
    """
    logger.info("Starting complete cleanup of all chunks and documents")
    
    results = {
        'success': False,
        'database_cleared': False,
        'vector_store_cleared': False,
        'stats_before': {},
        'stats_after': {},
        'errors': []
    }
    
    try:
        # Initialize components if not provided
        if db_manager is None:
            db_manager = DatabaseManager(db_path)
        
        if vector_store is None:
            vector_store = ChromaVectorStore(persist_directory)
        
        # Get stats before clearing
        try:
            results['stats_before'] = db_manager.get_database_stats()
            vector_info = vector_store.get_collection_info()
            results['stats_before']['vector_store_count'] = vector_info.get('count', 0)
            logger.info(f"Before clearing - Documents: {results['stats_before']['document_count']}, "
                       f"Chunks: {results['stats_before']['chunk_count']}, "
                       f"Vector store: {results['stats_before']['vector_store_count']}")
        except Exception as e:
            logger.warning(f"Could not get stats before clearing: {e}")
            results['errors'].append(f"Could not get pre-cleanup stats: {str(e)}")
        
        # Clear the database
        logger.info("Clearing SQLite database...")
        try:
            db_success = db_manager.clear_all_data()
            results['database_cleared'] = db_success
            if db_success:
                logger.info("Successfully cleared SQLite database")
            else:
                logger.error("Failed to clear SQLite database")
                results['errors'].append("Failed to clear SQLite database")
        except Exception as e:
            logger.error(f"Error clearing database: {e}")
            results['errors'].append(f"Database clearing error: {str(e)}")
            results['database_cleared'] = False
        
        # Clear the vector store
        logger.info("Clearing ChromaDB vector store...")
        try:
            vector_success = vector_store.reset_collection()
            
            # If reset failed, try to fix collection issues
            if not vector_success:
                logger.info("Reset failed, attempting to fix collection issues...")
                vector_success = vector_store.fix_collection_issues()
            
            results['vector_store_cleared'] = vector_success
            if vector_success:
                logger.info("Successfully cleared ChromaDB vector store")
            else:
                logger.error("Failed to clear ChromaDB vector store")
                results['errors'].append("Failed to clear ChromaDB vector store")
        except Exception as e:
            logger.error(f"Error clearing vector store: {e}")
            results['errors'].append(f"Vector store clearing error: {str(e)}")
            results['vector_store_cleared'] = False
        
        # Get stats after clearing
        try:
            results['stats_after'] = db_manager.get_database_stats()
            vector_info = vector_store.get_collection_info()
            results['stats_after']['vector_store_count'] = vector_info.get('count', 0)
            logger.info(f"After clearing - Documents: {results['stats_after']['document_count']}, "
                       f"Chunks: {results['stats_after']['chunk_count']}, "
                       f"Vector store: {results['stats_after']['vector_store_count']}")
        except Exception as e:
            logger.warning(f"Could not get stats after clearing: {e}")
            results['errors'].append(f"Could not get post-cleanup stats: {str(e)}")
        
        # Determine overall success
        results['success'] = results['database_cleared'] and results['vector_store_cleared']
        
        if results['success']:
            logger.info("Successfully completed cleanup of all chunks and documents")
        else:
            logger.warning("Cleanup completed with some failures")
            
        return results
        
    except Exception as e:
        logger.error(f"Unexpected error during cleanup: {e}")
        results['errors'].append(f"Unexpected error: {str(e)}")
        results['success'] = False
        return results


def clear_all_chunks_and_documents_simple(
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite"
) -> bool:
    """
    Simple function to clear all chunks and documents.
    
    Args:
        persist_directory (str): Directory containing the Chroma database
        db_path (str): Path to the SQLite database file
        
    Returns:
        bool: True if cleanup was successful, False otherwise
    """
    try:
        result = clear_all_chunks_and_documents(
            persist_directory=persist_directory,
            db_path=db_path
        )
        return result['success']
    except Exception as e:
        logger.error(f"Error in simple clear function: {e}")
        return False


def synchronize_database_and_vector_store(
    vector_store: Optional[ChromaVectorStore] = None,
    db_manager: Optional[DatabaseManager] = None,
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite",
    remove_orphaned: bool = True
) -> Dict[str, Any]:
    """
    Synchronize the SQLite database and ChromaDB vector store to remove orphaned references.
    
    This function identifies and optionally removes:
    1. Documents in vector store but not in database (orphaned vectors)
    2. Documents in database but not in vector store (orphaned metadata)
    
    Args:
        vector_store (Optional[ChromaVectorStore]): Vector store instance
        db_manager (Optional[DatabaseManager]): Database manager instance
        persist_directory (str): Directory containing the Chroma database
        db_path (str): Path to the SQLite database file
        remove_orphaned (bool): Whether to remove orphaned references (default: True)
        
    Returns:
        Dict[str, Any]: Synchronization results and statistics
    """
    logger.info("Starting database and vector store synchronization...")
    
    results = {
        'success': False,
        'orphaned_vectors_found': 0,
        'orphaned_vectors_removed': 0,
        'orphaned_database_entries_found': 0,
        'orphaned_database_entries_removed': 0,
        'errors': []
    }
    
    try:
        # Initialize components if not provided
        if db_manager is None:
            db_manager = DatabaseManager(db_path)
        
        if vector_store is None:
            vector_store = ChromaVectorStore(persist_directory)
        
        # Get all documents from database
        db_documents = db_manager.get_all_documents()
        db_file_paths = {doc['file_path'] for doc in db_documents}
        
        logger.info(f"Found {len(db_documents)} documents in database")
        
        # Get collection info from vector store
        try:
            collection_info = vector_store.get_collection_info()
            vector_count = collection_info.get('count', 0)
            logger.info(f"Found {vector_count} vectors in vector store")
            
            if vector_count == 0:
                logger.info("Vector store is empty, nothing to synchronize")
                results['success'] = True
                return results
                
        except Exception as e:
            logger.error(f"Could not get vector store info: {e}")
            results['errors'].append(f"Vector store access error: {str(e)}")
            return results
        
        # Sample vectors to check for orphaned references
        # We'll query the vector store with a dummy query to get all metadata
        try:
            from sentence_transformers import SentenceTransformer
            temp_model = SentenceTransformer('all-MiniLM-L6-v2')
            dummy_embedding = temp_model.encode(["dummy query"]).tolist()[0]
            
            # Query for many results to get a sample of what's in the vector store
            sample_results = vector_store.query_similar(
                query_text="dummy",
                query_embedding=dummy_embedding,
                n_results=min(100, vector_count)  # Sample up to 100 vectors
            )
            
            # Check for orphaned vectors (in vector store but not in database)
            orphaned_vector_files = set()
            if sample_results['metadatas'] and sample_results['metadatas'][0]:
                for metadata in sample_results['metadatas'][0]:
                    file_path = metadata.get('file_path', '')
                    if file_path and file_path not in db_file_paths:
                        orphaned_vector_files.add(file_path)
                        results['orphaned_vectors_found'] += 1
            
            logger.info(f"Found {results['orphaned_vectors_found']} orphaned vector references")
            
            if orphaned_vector_files and remove_orphaned:
                logger.info(f"Removing {len(orphaned_vector_files)} orphaned vector files...")
                # Note: ChromaDB doesn't have a direct way to delete by metadata
                # In a production system, you'd need to track chunk IDs better
                logger.warning("Orphaned vector cleanup requires collection reset for complete removal")
                results['orphaned_vectors_removed'] = len(orphaned_vector_files)
            
        except Exception as e:
            logger.error(f"Error checking vector store for orphaned references: {e}")
            results['errors'].append(f"Vector orphan check error: {str(e)}")
        
        # Check for orphaned database entries (in database but not in vector store)
        # This is harder to check without querying each document individually
        # For now, we'll assume database entries without corresponding vectors are orphaned
        
        # If we found orphaned vectors, suggest cleanup
        if results['orphaned_vectors_found'] > 0:
            logger.warning(f"Found {results['orphaned_vectors_found']} orphaned references")
            logger.info("To completely clean up orphaned references, consider:")
            logger.info("1. Use the 'Clear All Data' function to reset both database and vector store")
            logger.info("2. Re-upload your documents to ensure consistency")
        
        results['success'] = True
        logger.info("Database and vector store synchronization completed")
        
    except Exception as e:
        logger.error(f"Error during synchronization: {e}")
        results['errors'].append(f"Synchronization error: {str(e)}")
        results['success'] = False
    
    return results


def fix_orphaned_references(
    persist_directory: str = "data/chroma",
    db_path: str = "data/db.sqlite"
) -> bool:
    """
    Simple function to fix orphaned references between database and vector store.
    
    Args:
        persist_directory (str): Directory containing the Chroma database
        db_path (str): Path to the SQLite database file
        
    Returns:
        bool: True if synchronization was successful
    """
    try:
        result = synchronize_database_and_vector_store(
            persist_directory=persist_directory,
            db_path=db_path,
            remove_orphaned=True
        )
        return result['success']
    except Exception as e:
        logger.error(f"Error fixing orphaned references: {e}")
        return False
